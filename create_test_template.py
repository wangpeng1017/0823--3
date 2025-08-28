#!/usr/bin/env python3
"""
创建测试用的Word模板文件
包含各种类型的占位符用于测试替换功能
"""

from docx import Document
from docx.shared import Inches
import os

def create_test_template():
    """创建包含占位符的测试模板"""
    
    # 创建新文档
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('合同模板测试文档', 0)
    
    # 添加基本信息段落
    doc.add_heading('基本信息', level=1)
    
    basic_info = doc.add_paragraph()
    basic_info.add_run('甲方公司名称：{{甲方公司名称}}\n')
    basic_info.add_run('乙方公司名称：{{乙方公司名称}}\n')
    basic_info.add_run('合同编号：{{合同编号}}\n')
    basic_info.add_run('签署日期：{{签署日期}}\n')
    basic_info.add_run('合同金额：{{合同金额}}元\n')
    
    # 添加联系信息
    doc.add_heading('联系信息', level=1)
    
    contact_info = doc.add_paragraph()
    contact_info.add_run('甲方联系人：{{甲方联系人}}\n')
    contact_info.add_run('甲方电话：{{甲方电话}}\n')
    contact_info.add_run('甲方邮箱：{{甲方邮箱}}\n')
    contact_info.add_run('乙方联系人：{{乙方联系人}}\n')
    contact_info.add_run('乙方电话：{{乙方电话}}\n')
    
    # 添加合同条款
    doc.add_heading('合同条款', level=1)
    
    terms = doc.add_paragraph()
    terms.add_run('合同类型：{{合同类型}}\n')
    terms.add_run('服务内容：{{服务内容}}\n')
    terms.add_run('交付时间：{{交付时间}}\n')
    terms.add_run('付款方式：{{付款方式}}\n')
    
    # 添加备注
    doc.add_heading('备注', level=1)
    remarks = doc.add_paragraph()
    remarks.add_run('特殊说明：{{特殊说明}}\n')
    remarks.add_run('其他条款：{{其他条款}}\n')
    
    # 保存文档
    template_path = 'test-contract-template.docx'
    doc.save(template_path)
    
    print(f"✅ 测试模板已创建: {template_path}")
    print(f"📄 文件大小: {os.path.getsize(template_path)} bytes")
    
    # 显示包含的占位符
    placeholders = [
        '甲方公司名称', '乙方公司名称', '合同编号', '签署日期', '合同金额',
        '甲方联系人', '甲方电话', '甲方邮箱', '乙方联系人', '乙方电话',
        '合同类型', '服务内容', '交付时间', '付款方式', '特殊说明', '其他条款'
    ]
    
    print(f"🔖 包含占位符 ({len(placeholders)}个):")
    for i, placeholder in enumerate(placeholders, 1):
        print(f"  {i:2d}. {{{{ {placeholder} }}}}")
    
    return template_path

if __name__ == '__main__':
    try:
        template_path = create_test_template()
        print(f"\n🎯 模板创建成功！可以用于测试文档字段替换功能")
        print(f"📁 文件路径: {os.path.abspath(template_path)}")
    except Exception as e:
        print(f"❌ 创建模板失败: {e}")
        import traceback
        traceback.print_exc()
